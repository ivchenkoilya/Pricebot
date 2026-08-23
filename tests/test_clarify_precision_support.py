from __future__ import annotations

import pytest
from docx import Document

from app.ai.schemas import AnalysisResult
from app.bot.clarify_support import _looks_like_support_intent
from app.processors.common import retrieve_chunks
from app.processors.document_pipeline import extract_for_analysis, sanitize_document_analysis
from app.processors.documents import DocumentTooLarge


def _control_chunks() -> list[str]:
    return [
        '[Страница 4]\nСрок поставки считается соблюденным после приемки. Обычные условия поставки и документов.',
        '[Страница 37]\nЗа просрочку поставки начисляется неустойка 7,5% от стоимости просроченной партии за каждый полный календарный день, но не более 25% стоимости партии.',
        '[Страница 64]\nОкончательная оплата 4 850 000 рублей должна быть произведена не позднее 15.09.2026 после подписания акта приемки и получения корректного счета.',
        '[Страница 88]\nГарантийный срок 36 месяцев. Реакция сервиса не позднее 8 рабочих часов, выезд специалиста в течение 2 рабочих дней.',
        '[Страница 93]\nЗаказчик вправе расторгнуть договор в одностороннем порядке при трех и более нарушениях сроков поставки в течение 60 календарных дней. Поставщик должен вернуть полученный аванс в течение 5 банковских дней.',
    ]


def test_retrieval_finds_termination_and_refund_page_93():
    selected = retrieve_chunks(
        _control_chunks(),
        'Когда заказчик может расторгнуть договор и за сколько нужно вернуть аванс?',
        limit=3,
    )
    assert any('[Страница 93]' in chunk for chunk in selected)
    assert '[Страница 93]' in selected[-1] or '[Страница 93]' in selected[0]


def test_retrieval_finds_all_four_control_facts():
    cases = [
        ('Какой штраф за просрочку?', 'Страница 37'),
        ('Какая окончательная сумма оплаты и крайний срок?', 'Страница 64'),
        ('Какой гарантийный срок и срок реакции сервиса?', 'Страница 88'),
        ('Когда можно расторгнуть договор и вернуть аванс?', 'Страница 93'),
    ]
    chunks = _control_chunks()
    for query, marker in cases:
        selected = retrieve_chunks(chunks, query, limit=3)
        assert any(marker in chunk for chunk in selected), (query, selected)


def test_docx_explicit_page_breaks_are_real_pages(tmp_path):
    path = tmp_path / 'three-pages.docx'
    doc = Document()
    doc.add_paragraph('Первая страница')
    doc.add_page_break()
    doc.add_paragraph('Вторая страница')
    doc.add_page_break()
    doc.add_paragraph('Третья страница')
    doc.save(path)

    result = extract_for_analysis(str(path), '.docx', 3)

    assert result.pages == 3
    assert result.display_pages == '3 стр.'
    assert '[Страница 1]' in result.text
    assert '[Страница 2]' in result.text
    assert '[Страница 3]' in result.text


def test_docx_explicit_pages_enforce_plan_limit(tmp_path):
    path = tmp_path / 'three-pages.docx'
    doc = Document()
    doc.add_paragraph('1')
    doc.add_page_break()
    doc.add_paragraph('2')
    doc.add_page_break()
    doc.add_paragraph('3')
    doc.save(path)

    with pytest.raises(DocumentTooLarge, match='3 стр'):
        extract_for_analysis(str(path), '.docx', 2)


def test_clean_document_drops_invented_damage_warning():
    result = AnalysisResult(
        title='Документ',
        summary='Главное условие найдено. В документе есть обрезанные фрагменты текста.',
        warnings=['В документе есть поврежденные фрагменты текста.', 'Штраф 10%.'],
    )
    clean = sanitize_document_analysis(result)
    assert 'обрезан' not in clean.summary.lower()
    assert all('поврежден' not in item.lower() for item in clean.warnings)
    assert 'Штраф 10%.' in clean.warnings


def test_support_intent_is_conservative():
    assert _looks_like_support_intent('поддержка')
    assert _looks_like_support_intent('баг')
    assert _looks_like_support_intent('ошибка в боте')
    assert _looks_like_support_intent('Хочу написать в поддержку')
    assert _looks_like_support_intent('бот не работает')
    assert _looks_like_support_intent('сообщить об ошибке')
    assert not _looks_like_support_intent('Найди в этом документе раздел техническая поддержка')
    assert not _looks_like_support_intent('помоги понять договор')
