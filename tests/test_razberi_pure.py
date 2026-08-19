from pathlib import Path

from docx import Document
from openpyxl import Workbook

from app.processors.common import chunk_text, retrieve_chunks
from app.processors.documents import extract_docx, extract_text, extract_xlsx
from app.processors.router import InputKind, classify


def test_input_router():
    assert classify(text=True) == InputKind.TEXT
    assert classify(voice=True) == InputKind.VOICE
    assert classify(filename='report.pdf') == InputKind.PDF
    assert classify(filename='table.xlsx') == InputKind.XLSX
    assert classify(photo=True) == InputKind.IMAGE


def test_chunking_and_retrieval():
    text = ('Поставка оборудования состоится в пятницу. ' * 180) + 'Оплатить нужно 45000 рублей до четверга.'
    chunks = chunk_text(text, size=1500, overlap=100)
    assert len(chunks) > 2
    selected = retrieve_chunks(chunks, 'сколько оплатить до какого срока', limit=3)
    assert selected
    assert any('45000' in value for value in selected)


def test_real_document_extractors(tmp_path: Path):
    txt = tmp_path / 'a.txt'
    txt.write_text('Поставщик отправит заказ завтра.', encoding='utf-8')

    docx = tmp_path / 'a.docx'
    document = Document()
    document.add_heading('Договор', level=1)
    document.add_paragraph('Сумма 45000 рублей.')
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = 'Срок'
    table.cell(0, 1).text = 'Пятница'
    document.save(docx)

    xlsx = tmp_path / 'a.xlsx'
    wb = Workbook()
    ws = wb.active
    ws.append(['Товар', 'Цена'])
    ws.append(['A', 100])
    ws.append(['B', 250])
    wb.save(xlsx)

    assert 'Поставщик' in extract_text(str(txt))
    assert '45000' in extract_docx(str(docx))
    table_text = extract_xlsx(str(xlsx))
    assert '250' in table_text
    assert 'Цена' in table_text
