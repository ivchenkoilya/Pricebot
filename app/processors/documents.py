from __future__ import annotations

import csv
import io
from collections import defaultdict
from pathlib import Path
from statistics import mean

import fitz
from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook


class DocumentTooLarge(ValueError):
    pass


def extract_pdf(path: str, max_pages: int) -> tuple[str, int]:
    doc = fitz.open(path)
    try:
        pages = len(doc)
        if pages > max_pages:
            raise DocumentTooLarge(f'PDF: {pages} стр., лимит {max_pages}')
        parts: list[str] = []
        for index, page in enumerate(doc):
            text = page.get_text('text').strip()
            if text:
                parts.append(f'[Страница {index + 1}]\n{text}')
        return '\n\n'.join(parts), pages
    finally:
        doc.close()


def _explicit_page_break_count(paragraph) -> int:
    return sum(
        1
        for node in paragraph._p.iter(qn('w:br'))
        if node.get(qn('w:type')) == 'page'
    )


def _rendered_page_break_count(paragraph) -> int:
    return sum(1 for _ in paragraph._p.iter(qn('w:lastRenderedPageBreak')))


def extract_docx_with_page_info(path: str) -> tuple[str, int | None]:
    """Extract DOCX text and preserve the best page boundaries available.

    DOCX pagination normally depends on Word/LibreOffice rendering. We therefore
    prefer explicit <w:br type="page"> boundaries. If the document has no
    explicit breaks but Word stored lastRenderedPageBreak markers, those are used
    as the current rendered-page snapshot. We never add both kinds together, so
    one physical boundary cannot be double-counted.
    """
    doc = Document(path)
    paragraphs = list(doc.paragraphs)
    explicit_total = sum(_explicit_page_break_count(p) for p in paragraphs)
    rendered_total = sum(_rendered_page_break_count(p) for p in paragraphs)
    use_rendered = explicit_total == 0 and rendered_total > 0
    known_breaks = rendered_total if use_rendered else explicit_total

    parts: list[str] = ['[Страница 1]']
    page = 1
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
        breaks = (
            _rendered_page_break_count(paragraph)
            if use_rendered
            else _explicit_page_break_count(paragraph)
        )
        for _ in range(breaks):
            page += 1
            parts.append(f'[Страница {page}]')

    # python-docx exposes tables separately from paragraph flow. We keep them
    # after the body as before; exact table pagination would require rendering.
    for table_index, table in enumerate(doc.tables, 1):
        parts.append(f'[Таблица {table_index}]')
        for row in table.rows:
            parts.append(' | '.join(cell.text.strip() for cell in row.cells))

    known_pages = known_breaks + 1 if known_breaks else None
    return '\n'.join(parts), known_pages


def extract_docx(path: str) -> str:
    text, _pages = extract_docx_with_page_info(path)
    return text


def extract_text(path: str) -> str:
    return Path(path).read_text(encoding='utf-8', errors='replace')


def _numeric_stats(rows: list[list[object]]) -> list[str]:
    columns: dict[int, list[float]] = defaultdict(list)
    for row in rows:
        for index, value in enumerate(row):
            if isinstance(value, bool):
                continue
            if isinstance(value, (int, float)):
                columns[index].append(float(value))
    result: list[str] = []
    for index, values in sorted(columns.items()):
        if len(values) < 2:
            continue
        result.append(
            f'Колонка {index + 1}: n={len(values)}, min={min(values):g}, '
            f'max={max(values):g}, avg={mean(values):.2f}'
        )
    return result[:20]


def extract_xlsx(path: str, max_rows: int = 300) -> str:
    wb = load_workbook(path, read_only=True, data_only=True)
    output: list[str] = []
    for ws in wb.worksheets[:10]:
        output.append(f'[Лист: {ws.title}; строк: {ws.max_row}; колонок: {ws.max_column}]')
        captured: list[list[object]] = []
        for index, row in enumerate(ws.iter_rows(values_only=True)):
            values = list(row[:30])
            if index < max_rows:
                captured.append(values)
                output.append(' | '.join('' if value is None else str(value) for value in values))
            else:
                output.append(f'… остальные строки не переданы в AI после {max_rows}')
                break
        stats = _numeric_stats(captured)
        if stats:
            output.append('[Статистика числовых колонок]')
            output.extend(stats)
    return '\n'.join(output)


def extract_csv(path: str, max_rows: int = 300) -> str:
    text = Path(path).read_text(encoding='utf-8-sig', errors='replace')
    reader = csv.reader(io.StringIO(text))
    rows: list[list[str]] = []
    for index, row in enumerate(reader):
        if index >= max_rows:
            break
        rows.append(row[:30])
    output = [f'[CSV; показано строк: {len(rows)}]']
    output.extend(' | '.join(row) for row in rows)
    return '\n'.join(output)


def extract_document(path: str, ext: str, max_pages: int):
    ext = ext.lower()
    if ext == '.pdf':
        text, pages = extract_pdf(path, max_pages)
        return text, pages, 'pdf'
    if ext == '.docx':
        text, pages = extract_docx_with_page_info(path)
        if pages is not None and pages > max_pages:
            raise DocumentTooLarge(f'DOCX: {pages} стр., лимит {max_pages}')
        return text, pages, 'docx'
    if ext in {'.txt', '.md'}:
        return extract_text(path), None, 'text'
    if ext == '.xlsx':
        return extract_xlsx(path), None, 'spreadsheet'
    if ext == '.csv':
        return extract_csv(path), None, 'spreadsheet'
    raise ValueError('Неподдерживаемый формат')


def render_pdf_pages(path: str, max_pages: int = 10, dpi: int = 144):
    doc = fitz.open(path)
    output: list[tuple[int, bytes]] = []
    try:
        matrix = fitz.Matrix(dpi / 72, dpi / 72)
        for index, page in enumerate(doc):
            if index >= max_pages:
                break
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            output.append((index + 1, pix.tobytes('png')))
        return output
    finally:
        doc.close()
